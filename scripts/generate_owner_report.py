"""Generate owner follow-up Word report from git history and requirements tracker."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parent.parent / "CRM-Website-Progress-Report-v2.docx"


def set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_status_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ("Area", "Status", "Summary")
    header_cells = table.rows[0].cells
    for i, label in enumerate(headers):
        header_cells[i].text = label
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(header_cells[i], "1F4E79")
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    status_colors = {
        "Done": "C6EFCE",
        "Mostly done": "C6EFCE",
        "In progress": "FFEB9C",
        "Open": "FFC7CE",
        "Pending decision": "DDEBF7",
    }

    for req, status, notes in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = req
        row_cells[1].text = status
        row_cells[2].text = notes
        color = status_colors.get(status, "FFFFFF")
        set_cell_shading(row_cells[1], color)
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CRM & Website — Progress Report")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(
        f"Prepared for project owner | {date.today().strftime('%B %d, %Y')}"
    )
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()
    add_heading(doc, "Purpose", 2)
    doc.add_paragraph(
        "This report summarizes completed and remaining work on the Comma CRM and website, "
        "based on the original requirement document (CRM & Website Issues + New Feature Request) "
        "and verified against the frontend and backend git history (June 6–9, 2026)."
    )

    add_heading(doc, "Executive Summary", 2)
    doc.add_paragraph(
        "All six originally reported CRM issues have been fixed. The Projects lead-generation "
        "module is complete and ready for campaign use. Git repositories, CI/CD pipelines, and "
        "server structure cleanup are in place. The main remaining requirement from the PDF is "
        "the fully dynamic property attributes system (control-panel driven, not hardcoded)."
    )

    summary_rows = [
        ("Git repositories (frontend + backend)", "Done", "Both projects under commarealestate GitHub org"),
        ("CI/CD deployment pipelines", "Done", "Automated deploy to Hostinger on push to main"),
        ("Server structure & file cleanup", "Done", "Organized webroot; API in ~/apps/comma-api"),
        ("Existing CRM issues (6 items)", "Done", "All reported bugs resolved and verified"),
        ("Projects lead engine", "Done", "Landing pages, forms, CRM leads, homepage section"),
        ("CRM activity tracking", "Done", "Dates + timeline on agents and contact leads"),
        ("Dynamic property attributes", "Open", "Largest remaining PDF requirement"),
    ]
    add_status_table(doc, summary_rows)

    # --- Infrastructure ---
    add_heading(doc, "Infrastructure & DevOps (Completed)", 1)

    add_heading(doc, "Git Repositories", 2)
    git_items = [
        (
            "Frontend repository",
            "https://github.com/commarealestate/front-end — initial import committed Jun 6, 2026.",
        ),
        (
            "Backend repository",
            "https://github.com/commarealestate/backend — migrated from prior vendor repo; "
            "active origin set to commarealestate org.",
        ),
        (
            "Version control baseline",
            "Frontend codebase fully imported into git. Backend history preserved from initial "
            "Laravel commit (Feb 2026) through latest CRM fixes.",
        ),
    ]
    for title_text, detail in git_items:
        add_bullet(doc, f" — {detail}", bold_prefix=title_text)

    add_heading(doc, "CI/CD Pipelines", 2)
    cicd_items = [
        (
            "Frontend pipeline",
            "GitHub Actions workflow builds static Nuxt site (npm ci + generate), packages "
            "artifact, uploads to Hostinger via SSH, and publishes to commarealestate.ae webroot. "
            "Triggers on push to main and manual dispatch.",
        ),
        (
            "Backend pipeline",
            "GitHub Actions workflow SSH-deploys to Hostinger: git pull in ~/apps/comma-api, "
            "composer install, Laravel cache optimization, and migrate --force. "
            "Triggers on push to main and manual dispatch.",
        ),
        (
            "Build/deploy split",
            "Frontend build and deploy are separate jobs with artifact handoff for reliable releases.",
        ),
    ]
    for title_text, detail in cicd_items:
        add_bullet(doc, f" — {detail}", bold_prefix=title_text)

    add_heading(doc, "Server Structure & Cleanup", 2)
    server_items = [
        "Frontend deploy cleans the public webroot before each release, keeping only the apps folder and dated backup archives (api_old_*, public_html_old_*).",
        "Automatic tarball backup created before each frontend publish (~/backups/frontend-before-*.tar.gz).",
        "Backend API runs from a dedicated path on the server (~/apps/comma-api) instead of mixed public_html files.",
        "Property data cleanup commands added: deduplication, invalid import cleanup, and restore-from-backup tooling.",
        "Property sync command hardened for more reliable external API synchronization.",
        "API resource responses normalized (agent, blog, project, property, talent resources).",
    ]
    for item in server_items:
        add_bullet(doc, item)

    # --- Existing issues ---
    add_heading(doc, "Section 1 — Existing CRM Issues (All Completed)", 1)
    doc.add_paragraph(
        "These six items were reported in the original PDF. All are implemented and verified."
    )

    existing_issues = [
        (
            "1. Employee / Agent Creation",
            "Form no longer hangs. Validation errors display clearly. Duplicate email returns "
            "a field-level error (checks real e_mail column). Backend validation on create/update.",
        ),
        (
            "2. Property Images in Control Panel",
            "Admin property list and grid display images correctly. Normalized cover/gallery/floor-plan mapping.",
        ),
        (
            "3. Agent Profile Issues",
            "Full name (first + second/middle + last). Active/inactive status accurate. All fields "
            "persist after save/refresh. Profile image saves and displays. Work position supported.",
        ),
        (
            "4. Missing Dates / Activity Tracking",
            "Created, first contact, last activity, and last update dates on agents and leads. "
            "Full CRM activity timeline (Lead Created, WhatsApp Sent, Call Made, Meeting Scheduled, "
            "Reply Sent, Marked Read, Agent Created/Updated, etc.). Existing records backfilled.",
        ),
        (
            "5. Apartments Filter Not Showing",
            "Case/plural-tolerant filtering (apartment/Apartment/Apartments). Same for villas and townhouses.",
        ),
        (
            "6. Agent Page Levels & Sorting",
            "Three hierarchy groups on public site. Admin sets website_level and display_order. "
            "Sort by display priority, CRE ID, joining date.",
        ),
    ]
    for title_text, detail in existing_issues:
        add_bullet(doc, f" — {detail}", bold_prefix=title_text)

    # --- Projects ---
    add_heading(doc, "Section 2 — Projects Lead Generation Engine (Completed)", 1)
    projects_done = [
        "Projects module in website and control panel (not a blog — dedicated landing pages).",
        "Admin CRUD with Published / Draft status; backend stores, returns, and filters by status.",
        "Cover image + up to 5 gallery images with preview and removal before save.",
        "Homepage Projects section directly below hero (published projects only).",
        "Lead-first project landing page: contact form above the fold; details/gallery below.",
        "Per-project contact form (Name, Phone, Email, optional Message).",
        "Leads saved in CRM with project ID, title, source (Project Page), page URL, and UTM data.",
        "Admin contact table and search include project/campaign context.",
        "Ready for advertising campaigns — each project is a standalone conversion page.",
    ]
    for item in projects_done:
        add_bullet(doc, item)

    # --- CRM & Admin ---
    add_heading(doc, "Section 3 — CRM & Admin Improvements (Completed)", 1)
    crm_items = [
        "Agent admin form overhauled: full FormData submit, all mapped fields, image upload.",
        "Agent store rewritten for backend API CRUD with proper error handling.",
        "Contact message filters and search expanded (project, source, page URL, UTM).",
        "Contact admin table redesigned for project lead context and safer date formatting.",
        "Admin agent and contact detail pages show tracking dates and activity timeline.",
        "Requirements tracker (requirements-audit.md) rewritten in English with status per item.",
    ]
    for item in crm_items:
        add_bullet(doc, item)

    # --- Property & Website ---
    add_heading(doc, "Section 4 — Property & Website Improvements (Completed)", 1)
    property_items = [
        "Property detail page expanded with type-specific fields and contact form.",
        "Property leads include property metadata and listing agent context.",
        "Share action uses native share or clipboard copy (no browser alert).",
        "Arabic property location display fixed for Arabic locale pages.",
        "Property edit form Arabic location mapping fixed.",
        "Agent detail/back links preserve locale routing.",
        "Invalid undefined-* agent URLs fixed.",
        "Apartment filter normalization on frontend and backend.",
    ]
    for item in property_items:
        add_bullet(doc, item)

    # --- Backend data ---
    add_heading(doc, "Section 5 — Backend Data Quality (Completed)", 1)
    backend_items = [
        "Property sync command refactored for reliability.",
        "DedupeProperties command for removing duplicate listings.",
        "CleanupInvalidImportedProperties command for bad import data.",
        "RestorePropertiesFromDedupeBackup command for safe rollback.",
        "Database migration for property soft-deletes and xml_id uniqueness.",
        "Project status migration (add_status_to_projects_table).",
        "Agent migrations: display_order, website_level, second_name fields.",
        "Property migration: location_arabic field.",
        "Contact migration: project_id, project_title, source, page_url, UTM fields.",
    ]
    for item in backend_items:
        add_bullet(doc, item)

    # --- Git timeline ---
    add_heading(doc, "Development Timeline (from Git History)", 2)
    timeline_rows = [
        ("Jun 6", "Frontend", "Initial frontend import to git; backend property sync hardening & cleanup"),
        ("Jun 7", "Both", "CI/CD workflows added; agent form overhaul; project leads; contact admin; agent ordering"),
        ("Jun 7", "Frontend", "Property localization; project/property lead updates; build/deploy split"),
        ("Jun 8", "Both", "CRM verification pass; property contact/share improvements; backend project status"),
        ("Jun 9", "Both", "CRM activity tracking & timeline; requirements documentation update"),
    ]
    add_status_table(doc, [(d, r, n) for d, r, n in timeline_rows])

    # --- Remaining ---
    add_heading(doc, "Section 6 — Remaining Work", 1)
    doc.add_paragraph(
        "Only the following items from the original PDF requirements are not yet fully complete. "
        "Infrastructure, CI/CD, existing issues, and the Projects module are done."
    )

    remaining_rows = [
        (
            "Dynamic property attributes",
            "Open",
            "Resale, rental, and off-plan attributes must be managed from the control panel "
            "and rendered dynamically on the website. Currently property pages show many "
            "type-specific fields but they are not yet driven by a control-panel attribute schema.",
        ),
        (
            "Project rich-text editor",
            "Pending decision",
            "Project descriptions use a textarea today. Upgrade to a blog-style rich editor "
            "only if the owner requires it.",
        ),
        (
            "Draft project detail visibility",
            "Pending decision",
            "Published lists already filter correctly. A separate public/admin detail API "
            "split is optional if strict draft hiding on direct URLs is required.",
        ),
    ]
    add_status_table(doc, remaining_rows)

    add_heading(doc, "Verification Completed", 2)
    verification_done = [
        "Frontend production build passed (npm run build).",
        "Backend PHP syntax checks passed for all changed files.",
        "Project status and CRM tracking migrations ran successfully.",
        "Authenticated admin API verification passed for tracking fields and timeline.",
        "Nuxt dev server confirmed responding at localhost:3000.",
        "CI/CD workflows committed and configured with Hostinger deployment secrets.",
    ]
    for item in verification_done:
        add_bullet(doc, item)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note_run = note.add_run("Non-blocking notes: ")
    note_run.bold = True
    note.add_run(
        "Some frontend Sass/chunk-size build warnings remain. One unrelated default backend "
        "test fails (homepage route). Route listing blocked by an existing third-party payment "
        "package — not related to this work."
    )

    add_heading(doc, "Recommended Next Steps for Owner", 1)
    next_steps = [
        "Accept the six completed CRM issue fixes in staging or production.",
        "Test Projects end-to-end: create → publish → submit lead → verify in CRM with project/UTM data.",
        "Confirm priority and scope for dynamic property attributes (main remaining PDF item).",
        "Decide on rich-text editor for project descriptions (optional).",
        "Run UAT on campaign landing pages before go-live.",
    ]
    for step in next_steps:
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)

    add_heading(doc, "Reference", 2)
    doc.add_paragraph(
        "Source: CRM & Website Issues + New Feature Request (PDF). "
        "Repos: github.com/commarealestate/front-end, github.com/commarealestate/backend. "
        "Detailed tracker: requirements-audit.md."
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("— End of Report —")
    footer_run.italic = True
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Report saved to: {OUTPUT}")


if __name__ == "__main__":
    main()
